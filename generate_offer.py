from docx import Document
from transformers import GPT2LMHeadModel, GPT2Tokenizer 

# Incarcarea modelului pre-antrenat si tokenizer
model_name = 'gpt2'
model = GPT2LMHeadModel.from_pretrained(model_name)
tokenizer = GPT2Tokenizer.from_pretrained(model_name)

# Functia pentru generarea textului folosind modelul
def generate_text(prompt):
    inputs = tokenizer(prompt, return_tensors='pt', max_length=1024, truncation=True)
    outputs = model.generate(inputs['input_ids'], max_length=1024, num_return_sequences=1)
    text = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return text

# Functia pentru crearea ofertei in format DOCX
def create_offer(client_requirements, output_path):
    document = Document()

    # Adauga titlul documentului
    document.add_heading('Oferta pentru firma', 0)

    # Adauga solicitarea clientului
    document.add_heading('Solicitarea clientului', level=1)
    document.add_paragraph(client_requirements)

    # Adauga scopul documentului
    document.add_heading('Scopul documentului', level=1)
    prompt = "Scopul documentului: "
    document.add_paragraph(generate_text(prompt))

    # Adauga etapele de dezvoltare
    document.add_heading('Etapele de dezvoltare', level=1)
    stages = [
        "Elaborarea unei diagrame logice pentru a defini arhitectura aplicației.",
        "Crearea unei diagrame ER pentru a structura baza de date.",
        "Realizarea unui design inițial în Figma pentru a elimina orice ambiguitate legată de interfața grafică."
    ]
    for stage in stages:
        document.add_paragraph(stage)

    # Adauga costurile preliminare
    document.add_heading('Costuri preliminare', level=1)
    costs = [
        "Diagrama logică și diagrama ER: X Euro + TVA",
        "Proiectul în Figma: X Euro + TVA",
        "Acești bani se vor scădea din prețul total de dezvoltare odată acceptată oferta fermă."
    ]
    for cost in costs:
        document.add_paragraph(cost)

    # Adauga definitiile si tehnologiile
    document.add_heading('Definiții și tehnologii', level=1)
    technologies = [
        "React: React este o bibliotecă JavaScript pentru construirea interfețelor de utilizator. Este utilizată pentru crearea unor interfețe de utilizator reactive și eficiente din punct de vedere al performanței.",
        "Ionic: Ionic este un framework open-source pentru dezvoltarea de aplicații mobile hibride. Utilizează tehnologii web precum JavaScript/React/Angular pentru a construi aplicații pentru platforme mobile, cum ar fi iOS și Android."
    ]
    for tech in technologies:
        document.add_paragraph(tech)

    # Salveaza documentul
    document.save(output_path)

# Exemple de cerinte ale clientului
client_requirements = """
Aplicatia sa contina 3 ramuri esentiale: 
1. Sofer: harta cu traseul sau care sa contina punctele de ridicare si livrare. Timpul estimat pentru ambele. Posibilitatea de a prelua singur o solicitare de la un restaurant sau sa i se acorde livrari. Raport zilnic cu livrarile efectuate. Actualizare in timp real al parcursului sau catre client.
2. Restaurant: preluare comenzi, confirmarea lor si a statusului comenzii in timp real catre client. In orele de varf sa nu mai aiba clientul optiunea de a pune noi comenzi daca dureaza mai mult decat timpul estimat in aplicatie. Solictare sofer si optiunea de a vedea parcursul soferului.
3. Client: sa i-a la cunostinta toate detaliile legate de restaurant (timp estimate de livrare, timp de preparare etc). sa adauge produse in cos, achizitie cu card sau numerar din aplicatie. Sa primeasca notificari cu fiecare parcurs al comenzii (confirmata, comanda urmeaza sa fie ridicata, comanda a fost ridicata, comanda livrata). Dupa plasarea si confirmarea si preluarea comenzii de catre sofer sa I se prezinte soferul, numarul de contact al acestuia si parcursul sau pe harta.
"""

# Generarea si salvarea ofertei
output_path = 'generated_offer.docx'
create_offer(client_requirements, output_path)
